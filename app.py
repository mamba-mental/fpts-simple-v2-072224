# app.py
# Purpose: Flask application to handle file uploads and generate fundability reports.
# Last Modified: 2023-07-22
# Version: 1.6.0

## Version History:
# - v1.0.0: Initial version setup.
# - v1.1.0: Added JSON handling and dynamic values.
# - v1.2.0: Implemented gap analysis and fundability score calculation.
# - v1.3.0: Integrated Flask routes and file upload handling.
# - v1.4.0: Set up report rendering and conversion to PDF/DOCX.
# - v1.5.0: Added webhook handling for file uploads from ChatGPT.
# - v1.6.0: Removed hard-coded data, ensured webhook processes file URL and user data, improved error handling.

import os
import json
import requests
import pdfkit
from flask import Flask, request, jsonify
from docx import Document
from jinja2 import Template
from bs4 import BeautifulSoup
from datetime import datetime
from dateutil.relativedelta import relativedelta

app = Flask(__name__)

# Load JSON files
def load_json_file(file_path):
    with open(file_path, 'r') as file:
        return json.load(file)

# Calculate account age from year
def calculate_years_from_date(year, month=1):
    start_date = datetime(year, month, 1)
    today = datetime.today()
    return relativedelta(today, start_date).years

# Handle dynamic values in tradelines
def handle_dynamic_values(tradelines):
    for tradeline in tradelines:
        year = tradeline.get('Year', datetime.today().year)
        tradeline['Account_Age'] = calculate_years_from_date(year)
    return tradelines

business_primary_tradelines = handle_dynamic_values(load_json_file('data/business_primary_tradelines.json')['Standard_Tradeline_List'])
consumer_primary_tradelines = handle_dynamic_values(load_json_file('data/consumer_primary_tradelines.json')['Standard_Tradeline_List'])

# Extract text from uploaded file
def extract_text_from_file(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    # Add more conditions for other file types if needed
    return ""

# Extract client data from user data
def extract_client_data(user_data):
    return user_data  # Assuming user_data is already in the correct format

# Perform consumer gap analysis
def perform_consumer_gap_analysis(profile):
    # Example analysis logic
    return {
        "score": profile.get('consumer_average_fico_score', 0),
        "utilization": profile.get('credit_utilization', 0)
    }

# Perform business gap analysis
def perform_business_gap_analysis(profile):
    # Example analysis logic
    return {
        "score": profile.get('naics_code', 0),
        "business_age": profile.get('business_age', 0)
    }

# Calculate consumer fundability score
def calculate_consumer_fundability_score(profile):
    score = (
        profile.get('credit_utilization', 0) * 0.3 +
        profile.get('payment_history', 0) * 0.2 +
        profile.get('avg_account_age', 0) * 0.2 +
        profile.get('public_records', 0) * 0.1 +
        profile.get('new_credit_inquiries', 0) * 0.1 +
        profile.get('credit_mix', 0) * 0.1
    )
    capacity = score * 1750  # Example funding capacity
    return score, capacity

# Calculate business fundability score
def calculate_business_fundability_score(profile):
    score = (
        profile.get('naics_code', 0) * 0.25 +
        profile.get('credit_utilization', 0) * 0.25 +
        profile.get('business_age', 0) * 0.20 +
        profile.get('credit_mix', 0) * 0.10 +
        profile.get('consumer_average_fico_score', 0) * 0.20
    )
    capacity = score * 5000  # Business funding capacity might be higher
    return score, capacity

# Recommend tradelines
def recommend_tradelines(profile, tradeline_list, min_recommendations=2):
    return tradeline_list[:min_recommendations]

# Generate consumer recommendations
def generate_consumer_recommendations(profile, tradeline_list, au_list):
    good_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=2)
    better_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=4)
    best_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=6)
    
    au_recommendations = au_list[:2]  # Example

    good_score, good_capacity = calculate_consumer_fundability_score(profile)
    better_score, better_capacity = good_score * 1.5, good_capacity * 1.5
    best_score, best_capacity = good_score * 1.75, good_capacity * 1.75

    recommendations = {
        'Good': {
            'Title': 'Good: Essential Improvements',
            'Description': 'Basic improvements that address the most critical gaps with minimal investment.',
            'Tradelines': good_recommendations,
            'AUs': au_recommendations,
            'Fundability_Score': good_score,
            'Estimated_Funding_Capacity': good_capacity
        },
        'Better': {
            'Title': 'Better: Enhanced Improvements',
            'Description': 'Building on essential improvements for a more robust profile.',
            'Tradelines': better_recommendations,
            'AUs': au_recommendations,
            'Fundability_Score': better_score,
            'Estimated_Funding_Capacity': better_capacity
        },
        'Best': {
            'Title': 'Best: Maximum Improvements',
            'Description': 'Comprehensive improvements for maximum funding potential.',
            'Tradelines': best_recommendations,
            'AUs': au_recommendations,
            'Fundability_Score': best_score,
            'Estimated_Funding_Capacity': best_capacity
        }
    }
    return recommendations

# Generate business recommendations
def generate_business_recommendations(profile, tradeline_list):
    good_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=2)
    better_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=4)
    best_recommendations = recommend_tradelines(profile, tradeline_list, min_recommendations=6)

    good_score, good_capacity = calculate_business_fundability_score(profile)
    better_score, better_capacity = good_score * 1.5, good_capacity * 1.5
    best_score, best_capacity = good_score * 1.75, good_capacity * 1.75

    recommendations = {
        'Good': {
            'Title': 'Good: Essential Improvements',
            'Description': 'Basic improvements that address the most critical gaps with minimal investment.',
            'Tradelines': good_recommendations,
            'Fundability_Score': good_score,
            'Estimated_Funding_Capacity': good_capacity
        },
        'Better': {
            'Title': 'Better: Enhanced Improvements',
            'Description': 'Building on essential improvements for a more robust profile.',
            'Tradelines': better_recommendations,
            'Fundability_Score': better_score,
            'Estimated_Funding_Capacity': better_capacity
        },
        'Best': {
            'Title': 'Best: Maximum Improvements',
            'Description': 'Comprehensive improvements for maximum funding potential.',
            'Tradelines': best_recommendations,
            'Fundability_Score': best_score,
            'Estimated_Funding_Capacity': best_capacity
        }
    }
    return recommendations

# Render the report
def render_report(client_profile, consumer_gap_analysis, business_gap_analysis, consumer_recommendations, business_recommendations):
    html_template = open('templates/report_template.html').read()
    template = Template(html_template)
    rendered_html = template.render(
        profile=client_profile,
        creator_name="Your Name",
        report_date=datetime.today().strftime('%Y-%m-%d'),
        introduction_summary="This report provides a comprehensive evaluation...",
        consumer_gap_analysis=consumer_gap_analysis,
        business_gap_analysis=business_gap_analysis,
        consumer_recommendations=consumer_recommendations,
        business_recommendations=business_recommendations
    )

    with open('output/funding_report.html', 'w') as file:
        file.write(rendered_html)

    pdfkit.from_file('output/funding_report.html', 'output/funding_report.pdf')

    html_content = open('output/funding_report.html').read()
    soup = BeautifulSoup(html_content, 'html.parser')

    doc = Document()
    doc.add_heading('Funding Potential & Tradeline Strategy Report', 0)

    for element in soup.body.children:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'table':
            table = doc.add_table(rows=1, cols=len(element.find_all('th')))
            hdr_cells = table.rows[0].cells
            for i, th in enumerate(element.find_all('th')):
                hdr_cells[i].text = th.text
            for row in element.find_all('tr')[1:]:
                row_cells = table.add_row().cells
                for i, td in enumerate(row.find_all('td')):
                    row_cells[i].text = td.text

    doc.save('output/funding_report.docx')
    return rendered_html

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"})
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"})
    if file:
        file_path = os.path.join("uploads", file.filename)
        file.save(file_path)
        text_data = extract_text_from_file(file_path)
        client_profile = extract_client_data(user_data)
        consumer_gap_analysis = perform_consumer_gap_analysis(client_profile)
        business_gap_analysis = perform_business_gap_analysis(client_profile)
        consumer_recommendations = generate_consumer_recommendations(client_profile, consumer_primary_tradelines, au_data)
        business_recommendations = generate_business_recommendations(client_profile, business_primary_tradelines)
        report = render_report(client_profile, consumer_gap_analysis, business_gap_analysis, consumer_recommendations, business_recommendations)
        return jsonify({"message": "Data extraction complete", "report": report})
    return jsonify({"error": "File upload failed"})

@app.route('/webhook', methods=['POST'])
def webhook():
    data = request.json
    file_url = data.get('file_url')
    user_data = data.get('user_data')
    
    if not file_url or not user_data:
        return jsonify({"error": "File URL or user data not provided"}), 400
    
    response = requests.get(file_url)
    if response.status_code != 200:
        return jsonify({"error": "Failed to download the file"}), 400
    
    file_path = os.path.join("uploads", "client_evaluation.docx")
    with open(file_path, 'wb') as file:
        file.write(response.content)
    
    text_data = extract_text_from_file(file_path)
    client_profile = extract_client_data(user_data)
    consumer_gap_analysis = perform_consumer_gap_analysis(client_profile)
    business_gap_analysis = perform_business_gap_analysis(client_profile)
    consumer_recommendations = generate_consumer_recommendations(client_profile, consumer_primary_tradelines, au_data)
    business_recommendations = generate_business_recommendations(client_profile, business_primary_tradelines)
    report = render_report(client_profile, consumer_gap_analysis, business_gap_analysis, consumer_recommendations, business_recommendations)
    
    return jsonify({"message": "Report generated", "report_url": "output/funding_report.pdf"})

@app.route('/generate_report', methods=['GET'])
def generate_report():
    return jsonify({"message": "Report generation initiated"})

if __name__ == '__main__':
    app.run(debug=True)
# End of file